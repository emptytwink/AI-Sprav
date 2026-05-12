from __future__ import annotations

import csv
import html
import json
import os
import re
from io import StringIO
from pathlib import Path
from typing import Any, Dict, List, Tuple

try:
    from dotenv import load_dotenv
except Exception:  # pragma: no cover
    def load_dotenv() -> bool:
        return False

try:
    from bs4 import BeautifulSoup
except Exception:  # pragma: no cover
    BeautifulSoup = None

SUPPORTED_EXTS = {
    ".pdf", ".docx", ".xlsx", ".xlsm",
    ".txt", ".md", ".html", ".htm",
    ".csv", ".json", ".xml", ".yaml", ".yml", ".log", ".ini", ".conf",
}
TRASH_NAMES = {".DS_Store", "Thumbs.db", "desktop.ini"}
TEXT_EXTS = {".txt", ".md", ".log", ".ini", ".conf", ".xml", ".yaml", ".yml"}


def collect_documents(source_dir: Path) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]], Dict[str, Any]]:
    load_dotenv()
    max_docs = int(os.getenv("AI_MAX_DOCS", "50"))
    max_chars_per_doc = int(os.getenv("AI_MAX_CHARS_PER_DOC", "4000"))

    all_files = [p for p in source_dir.rglob("*") if p.is_file() and not _is_trash(p, source_dir)]
    supported = [p for p in all_files if p.suffix.lower() in SUPPORTED_EXTS and "control_expected_result" not in p.relative_to(source_dir).parts]
    unsupported = [p for p in all_files if p not in supported]
    folders = sorted({_folder_path(p, source_dir) for p in all_files if _folder_path(p, source_dir)})

    selected = _select_evenly(supported, source_dir, max_docs)
    documents: List[Dict[str, Any]] = []
    skipped: List[Dict[str, Any]] = []

    for path in selected:
        rel = path.relative_to(source_dir).as_posix()
        folder_path, folder_parts = _folder_info(path, source_dir)
        text, encoding, error = extract_text(path)
        ext = path.suffix.lower()
        if error:
            skipped.append({
                "path": rel,
                "relpath": rel,
                "folder_path": folder_path,
                "folder_parts": folder_parts,
                "depth": len(folder_parts),
                "ext": ext,
                "reason": error,
                "skipped_reason": error,
            })
            continue
        text = " ".join(text.split())
        if not text:
            reason = "текст не найден"
            skipped.append({
                "path": rel,
                "relpath": rel,
                "folder_path": folder_path,
                "folder_parts": folder_parts,
                "depth": len(folder_parts),
                "ext": ext,
                "reason": reason,
                "skipped_reason": reason,
            })
            continue
        documents.append(
            {
                "path": rel,
                "relpath": rel,
                "folder_path": folder_path,
                "folder_parts": folder_parts,
                "depth": len(folder_parts),
                "title": path.stem,
                "ext": ext,
                "encoding": encoding,
                "chars_count": len(text),
                "text": text[:max_chars_per_doc],
                "extracted_text": text[:max_chars_per_doc],
            }
        )

    selected_set = set(selected)
    for path in supported:
        if path not in selected_set:
            rel = path.relative_to(source_dir).as_posix()
            folder_path, folder_parts = _folder_info(path, source_dir)
            reason = f"не выбран из-за ограничения AI_MAX_DOCS={max_docs}"
            skipped.append(
                {
                    "path": rel,
                    "relpath": rel,
                    "folder_path": folder_path,
                    "folder_parts": folder_parts,
                    "depth": len(folder_parts),
                    "ext": path.suffix.lower(),
                    "reason": reason,
                    "skipped_reason": reason,
                }
            )

    for path in unsupported:
        rel = path.relative_to(source_dir).as_posix()
        folder_path, folder_parts = _folder_info(path, source_dir)
        reason = "Формат не используется для текстового анализа"
        skipped.append(
            {
                "path": rel,
                "relpath": rel,
                "folder_path": folder_path,
                "folder_parts": folder_parts,
                "depth": len(folder_parts),
                "ext": path.suffix.lower(),
                "reason": reason,
                "skipped_reason": reason,
            }
        )

    max_archive_depth = max((len(path.relative_to(source_dir).parts) - 1 for path in all_files), default=0)
    meta = {
        "total_files_in_archive": len(all_files),
        "total_supported_candidates": len(supported),
        "selected_documents_count": len(documents),
        "selected_documents": [{k: v for k, v in doc.items() if k not in {"text", "extracted_text"}} for doc in documents],
        "skipped_files_count": len(skipped),
        "skipped_files": skipped,
        "folders_detected": folders,
        "max_archive_depth": max_archive_depth,
        "max_docs": max_docs,
        "max_chars_per_doc": max_chars_per_doc,
        "num_predict": int(os.getenv("AI_NUM_PREDICT", "8192")),
    }
    return documents, skipped, meta


def extract_text(path: Path) -> Tuple[str, str | None, str | None]:
    ext = path.suffix.lower()
    try:
        if ext in TEXT_EXTS:
            text, enc = read_text_safely(path)
            return text, enc, None
        if ext in {".html", ".htm"}:
            text, enc = read_text_safely(path)
            return _extract_html(text), enc, None
        if ext == ".csv":
            text, enc = read_text_safely(path)
            return _extract_csv(text), enc, None
        if ext == ".json":
            text, enc = read_text_safely(path)
            return _extract_json(text), enc, None
        if ext == ".pdf":
            text, error = _extract_pdf(path)
            return text, None, error
        if ext == ".docx":
            text, error = _extract_docx(path)
            return text, None, error
        if ext in {".xlsx", ".xlsm"}:
            text, error = _extract_xlsx(path)
            return text, None, error
    except Exception as exc:
        return "", None, f"ошибка чтения: {exc}"
    return "", None, "формат не поддерживается"


def read_text_safely(path: Path) -> Tuple[str, str]:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1251", "latin-1"):
        try:
            return data.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    return data.decode("latin-1", errors="replace"), "latin-1"


def _extract_html(raw: str) -> str:
    if BeautifulSoup is None:
        return re.sub(r"<[^>]+>", " ", html.unescape(raw))
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text(" ", strip=True)


def _extract_csv(raw: str) -> str:
    reader = csv.reader(StringIO(raw))
    lines = []
    for row in reader:
        lines.append(" | ".join(cell.strip() for cell in row))
        if len(lines) >= 200:
            break
    return "\n".join(lines)


def _extract_json(raw: str) -> str:
    try:
        return json.dumps(json.loads(raw), ensure_ascii=False, indent=2)
    except Exception:
        return raw


def _extract_pdf(path: Path) -> Tuple[str, str | None]:
    try:
        import fitz
    except Exception:
        return "", "PyMuPDF не установлен"

    max_chars_per_doc = int(os.getenv("AI_MAX_CHARS_PER_DOC", "4000"))
    chunks: List[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            chunks.append(page.get_text("text"))
            if sum(len(x) for x in chunks) >= max_chars_per_doc:
                break
    return "\n".join(chunks), None


def _extract_docx(path: Path) -> Tuple[str, str | None]:
    try:
        import docx
    except Exception:
        return "", "python-docx не установлен"

    document = docx.Document(path)
    chunks = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n".join(chunks), None


def _extract_xlsx(path: Path) -> Tuple[str, str | None]:
    try:
        from openpyxl import load_workbook
    except Exception:
        return "", "openpyxl не установлен"

    max_chars_per_doc = int(os.getenv("AI_MAX_CHARS_PER_DOC", "4000"))
    max_rows_per_sheet = int(os.getenv("AI_XLSX_MAX_ROWS_PER_SHEET", "200"))
    chunks: List[str] = []
    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        for sheet in workbook.worksheets:
            chunks.append(f"Лист: {sheet.title}")
            rows_read = 0
            for row in sheet.iter_rows(values_only=True):
                values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if not values:
                    continue
                chunks.append(" | ".join(values))
                rows_read += 1
                if rows_read >= max_rows_per_sheet:
                    break
                if sum(len(x) for x in chunks) >= max_chars_per_doc:
                    return "\n".join(chunks), None
    finally:
        workbook.close()
    return "\n".join(chunks), None


def _select_evenly(paths: List[Path], source_dir: Path, max_docs: int) -> List[Path]:
    """Берём документы равномерно по глубоким веткам, а не только по первой папке."""
    if max_docs <= 0:
        return []

    groups: Dict[str, List[Path]] = {}
    for path in paths:
        groups.setdefault(_selection_key(path, source_dir), []).append(path)
    for items in groups.values():
        items.sort(key=lambda p: (_priority(p), len(p.relative_to(source_dir).parts), str(p).lower()))

    selected: List[Path] = []

    # Первый проход: минимум по одному документу из каждой глубокой ветки.
    for key in sorted(groups):
        if len(selected) >= max_docs:
            return selected
        if groups[key]:
            selected.append(groups[key][0])

    # Второй проход: ещё 1-2 документа из каждой ветки, если лимит позволяет.
    for round_index in (1, 2):
        for key in sorted(groups):
            if len(selected) >= max_docs:
                return selected
            if len(groups[key]) > round_index:
                selected.append(groups[key][round_index])

    selected_set = set(selected)
    remaining = [p for p in paths if p not in selected_set]
    remaining.sort(key=lambda p: (_priority(p), len(p.relative_to(source_dir).parts), str(p).lower()))
    for path in remaining:
        if len(selected) >= max_docs:
            break
        selected.append(path)
    return selected


def _priority(path: Path) -> int:
    order = {
        ".pdf": 0,
        ".docx": 1,
        ".xlsx": 2,
        ".xlsm": 2,
        ".html": 3,
        ".htm": 3,
        ".csv": 4,
        ".json": 5,
        ".txt": 6,
        ".md": 7,
        ".xml": 8,
        ".yaml": 9,
        ".yml": 9,
        ".log": 10,
        ".ini": 11,
        ".conf": 11,
    }
    return order.get(path.suffix.lower(), 99)


def _selection_key(path: Path, source_dir: Path) -> str:
    rel = path.relative_to(source_dir)
    # Берём первые 3 папки, чтобы охватить глубокую структуру архива.
    folder_parts = rel.parts[:-1]
    if not folder_parts:
        return "."
    return "/".join(folder_parts[:3])


def _folder_path(path: Path, source_dir: Path) -> str:
    return "/".join(path.relative_to(source_dir).parts[:-1])


def _folder_info(path: Path, source_dir: Path) -> Tuple[str, List[str]]:
    folder_parts = list(path.relative_to(source_dir).parts[:-1])
    return "/".join(folder_parts), folder_parts


def _is_trash(path: Path, source_dir: Path) -> bool:
    rel = path.relative_to(source_dir)
    return (
        any(part == "__MACOSX" for part in rel.parts)
        or path.name in TRASH_NAMES
        or "control_expected_result" in rel.parts
    )
